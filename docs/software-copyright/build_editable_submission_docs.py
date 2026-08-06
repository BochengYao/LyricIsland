from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT / "output" / "软件著作权登记材料-V2.0.36"
SOFTWARE_NAME = "桌面歌词岛软件"
SHORT_NAME = "歌词岛"
VERSION = "V2.0.36"
BUILD_VERSION = "2.0.36-Beta"
OWNER = "么博丞"


def count_source_lines() -> int:
    total = 0
    for project in (ROOT / "LyricsIsland.App", ROOT / "LyricsIsland.Core"):
        for path in list(project.rglob("*.cs")) + list(project.rglob("*.xaml")):
            if "bin" in path.parts or "obj" in path.parts:
                continue
            total += sum(
                1
                for line in path.read_text(encoding="utf-8-sig").splitlines()
                if line.strip()
            )
    return total


SOURCE_LINES = count_source_lines()

BLUE = "2667C9"
INK = "152238"
MUTED = "657083"
LIGHT_BLUE = "EAF2FF"
LIGHT_GRAY = "F2F4F7"
RULE = "CBD3DF"


def set_run_font(run, name: str = "Microsoft YaHei", size: float = 10.5,
                 bold: bool | None = None, color: str = INK) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, start: int = 120,
                     bottom: int = 90, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Cm(width / 1440 * 2.54)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_field(paragraph) -> None:
    paragraph.add_run("第 ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    paragraph.add_run(" 页")


def configure_document(doc: Document, running_label: str) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in (
        ("Title", 24, INK, 0, 8),
        ("Subtitle", 12, MUTED, 0, 16),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, INK, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(running_label), size=8, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    add_page_field(p)
    for run in p.runs:
        set_run_font(run, size=8, color=MUTED)


def add_title_block(doc: Document, title: str, subtitle: str, status: str) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(p.add_run(title), size=24, bold=True)
    p = doc.add_paragraph()
    p.style = doc.styles["Subtitle"]
    set_run_font(p.add_run(subtitle), size=12, color=MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    set_run_font(p.add_run(status), size=9.5, bold=True, color=BLUE)


def add_info_table(doc: Document, rows: list[tuple[str, str]], widths=(2500, 6860)) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = ""
        cells[1].text = ""
        set_cell_shading(cells[0], LIGHT_GRAY)
        lp = cells[0].paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        set_run_font(lp.add_run(label), size=9.5, bold=True)
        vp = cells[1].paragraphs[0]
        vp.paragraph_format.space_after = Pt(0)
        set_run_font(vp.add_run(value), size=9.5)
    set_table_geometry(table, list(widths))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_matrix(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(text), size=9, bold=True, color=BLUE)
    set_repeat_table_header(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(values):
            cells[idx].text = ""
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, len(values) - 1) else WD_ALIGN_PARAGRAPH.LEFT
            set_run_font(p.add_run(text), size=8.8)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_note(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_cell_shading(table.cell(0, 0), LIGHT_BLUE)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(f"{label}："), size=9.5, bold=True, color=BLUE)
    set_run_font(p.add_run(text), size=9.5)
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def build_application_brief() -> Path:
    doc = Document()
    configure_document(doc, f"{SOFTWARE_NAME} {VERSION}｜软著申报填报底稿")
    add_title_block(
        doc,
        "计算机软件著作权登记申报填报底稿",
        f"{SOFTWARE_NAME} {VERSION}（对应当前构建 {BUILD_VERSION}）",
        "用途：在中国版权保护中心登记系统逐项填写前核对；本文件本身不是官方申请表。",
    )

    doc.add_heading("1. 建议登记信息", level=1)
    add_info_table(doc, [
        ("软件全称", SOFTWARE_NAME),
        ("软件简称", SHORT_NAME),
        ("申报版本号", VERSION),
        ("当前构建标识", BUILD_VERSION),
        ("软件分类", "应用软件"),
        ("开发方式", "独立开发"),
        ("权利取得方式", "原始取得"),
        ("权利范围", "全部权利"),
        ("申请人／著作权人", OWNER),
        ("开发语言", "C#、XAML"),
        ("源程序量", f"当前自有 C#/XAML 源程序约 {SOURCE_LINES:,} 非空行"),
    ])
    add_note(
        doc,
        "版本一致性",
        f"正式申请表、源程序、说明书和补充材料统一填写 {SOFTWARE_NAME} {VERSION}。"
        f"“{BUILD_VERSION}”只用于说明其与当前 Beta 构建的对应关系，不作为申报版本号。",
    )

    doc.add_heading("2. 开发与运行环境", level=1)
    add_info_table(doc, [
        ("开发硬件环境", "x64 计算机；建议 8 GB 及以上内存；常规磁盘空间"),
        ("开发软件环境", "Windows 10/11、Visual Studio、.NET SDK、Windows 10 SDK"),
        ("运行硬件环境", "x64 处理器；建议 4 GB 及以上内存；支持桌面显示器"),
        ("运行软件环境", "Windows 10/11 x64；Microsoft Windows Desktop Runtime 3.1"),
        ("开发工具", "Visual Studio、Git、.NET CLI"),
        ("主要技术", "WPF、Windows SMTC、异步网络请求、本地 JSON 配置与文件缓存"),
    ])

    doc.add_heading("3. 主要功能与技术特点", level=1)
    add_info_table(doc, [
        ("开发目的", "为 Windows 用户提供不依赖单一音乐客户端的桌面同步歌词、播放信息与轻量控制能力。"),
        ("面向领域", "桌面工具、数字音乐辅助、媒体信息展示。"),
        ("主要功能", "读取 Windows 系统媒体会话；自动选择或锁定播放器；从多个歌词源匹配同步歌词及已有翻译；以顶部歌词岛显示歌词、封面、歌曲信息、播放控制和进度；支持真实悬浮窗口模块拖放编辑、快捷键校准、多显示器定位、鼠标避让和本地缓存。"),
        ("技术特点", "统一媒体快照与会话选择；时间轴可靠性判断及单调时钟补偿；多歌词源候选匹配；模块化布局投影；真实歌词岛跨窗口拖放；草稿保存与取消回滚；按能力逐项降级。"),
    ])

    doc.add_heading("4. 必须由申请人确认的事实", level=1)
    add_matrix(
        doc,
        ["项目", "需要填写或确认", "当前状态"],
        [
            ["身份信息", "身份证号、证件有效期、手机号、邮箱、通信地址、邮编", "待本人填写"],
            ["完成日期", "软件实际开发完成日期，须与事实和证据一致", "待本人确认"],
            ["发表状态", "按 V2.0.36 是否已向公众提供下载或访问的事实选择", "待本人确认"],
            ["发表信息", "若已发表，填写真实首次发表日期和首次发表地点", "条件填写"],
            ["权属事实", "确认代码为本人独立开发；如有合作、委托或职务开发须按事实调整", "待本人确认"],
            ["签章", "按登记系统生成的确认文件要求签名，姓名与身份证一致", "提交前完成"],
        ],
        [1450, 6200, 1710],
    )
    add_note(
        doc,
        "不要代填",
        "开发完成日期、发表状态和权属事实会影响申请内容，无法仅从仓库可靠推定。本底稿故意不替申请人作法律事实判断。",
    )

    doc.add_heading("5. 已准备的提交材料", level=1)
    add_matrix(
        doc,
        ["材料", "内容与规格", "处理建议"],
        [
            ["源程序鉴别材料", "60 页；每页 50 行；前、后各连续 30 页；A4；页眉含全称与版本", "上传 PDF"],
            ["软件说明书", "软件概述、环境、架构、功能、操作流程、异常处理和权利声明", "上传 PDF"],
            ["填报底稿", "登记系统各字段建议值、待确认事实和提交检查项", "内部核对"],
            ["新增版本说明", "说明 V2.0.36 相比早期版本的主要功能与技术变化", "系统要求时上传"],
            ["身份证明", "申请人身份证正反面或系统要求的实名材料", "本人准备"],
            ["申请确认文件", "登记系统生成后按要求签名或盖章", "提交前完成"],
        ],
        [1900, 5200, 2260],
    )

    doc.add_heading("6. 提交前一致性检查", level=1)
    add_matrix(
        doc,
        ["检查项", "通过标准", "结果"],
        [
            ["名称", f"所有材料均为“{SOFTWARE_NAME}”", "□"],
            ["版本", f"所有申报材料均为“{VERSION}”", "□"],
            ["权利人", f"所有材料均为“{OWNER}”且签名一致", "□"],
            ["页数", "源程序 60 页；说明书少于 60 页时提交全部", "□"],
            ["页眉页码", "源程序和说明书均可识别软件名称、版本和连续页码", "□"],
            ["隐私", "上传前确认源程序样本不含密钥、令牌、个人账号或真实用户数据", "□"],
            ["事实字段", "完成日期、发表状态、发表日期地点均已按事实填写", "□"],
            ["文件可读性", "PDF 可打开、字体正常、无裁切、无空白页或重复页", "□"],
        ],
        [2200, 6260, 900],
    )

    doc.add_heading("7. 办理入口与规则依据", level=1)
    p = doc.add_paragraph()
    set_run_font(p.add_run("办理入口：中国版权保护中心官网 https://www.ccopyright.com.cn/ 。"), size=10)
    p = doc.add_paragraph()
    set_run_font(
        p.add_run(
            "规则核对：国家版权局《计算机软件著作权登记办法》第九至十七条。"
            "其中规定申请表、程序和文档鉴别材料及相关证明文件的基本构成，"
            "并规定 A4 纸张、程序和文档前后各连续 30 页等要求。"
        ),
        size=10,
    )
    p = doc.add_paragraph()
    set_run_font(
        p.add_run(
            "提示：登记系统页面、实名认证流程、上传格式和补正要求可能调整，"
            "最终以提交当日中国版权保护中心系统提示为准。"
        ),
        size=10,
        color=MUTED,
    )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / f"{SOFTWARE_NAME}{VERSION}-申报信息填报底稿.docx"
    doc.save(path)
    return path


def build_version_statement() -> Path:
    doc = Document()
    configure_document(doc, f"{SOFTWARE_NAME} {VERSION}｜新增版本说明")
    add_title_block(
        doc,
        "新增版本说明",
        f"{SOFTWARE_NAME} {VERSION}",
        "用途：当登记系统或补正通知要求说明高版本软件的新增功能时提交。",
    )

    doc.add_heading("一、软件基本信息", level=1)
    add_info_table(doc, [
        ("软件全称", SOFTWARE_NAME),
        ("软件简称", SHORT_NAME),
        ("申报版本号", VERSION),
        ("对应构建", BUILD_VERSION),
        ("著作权人", OWNER),
        ("开发方式", "独立开发"),
        ("权利取得方式", "原始取得"),
    ])

    doc.add_heading("二、本版本主要新增与完善内容", level=1)
    add_matrix(
        doc,
        ["序号", "功能或技术变化", "具体说明"],
        [
            ["1", "多播放器媒体会话", "由单一音乐客户端场景扩展为读取 Windows SMTC，可自动跟随最近活跃会话或锁定指定播放器。"],
            ["2", "模块化歌词岛", "将封面、歌词、歌曲信息、播放控制、进度和分割线拆分为可组合模块，并支持不同布局模式。"],
            ["3", "真实界面布局编辑", "可从设置窗口把模块拖到正在运行的真实歌词岛，支持吸附、重排、草稿预览、保存和取消回滚。"],
            ["4", "歌词匹配与翻译", "支持多个歌词来源的候选检索、匹配、同步时间轴和歌词库已有逐行翻译，并提供首选来源和缓存。"],
            ["5", "时间轴补偿", "判断播放器时间轴可靠性；缺失或停滞时使用单调时钟估算，恢复可信时间轴后完成校准。"],
            ["6", "播放控制与进度", "按媒体会话实际能力显示上一曲、播放/暂停、下一曲和进度，并对不支持能力独立降级。"],
            ["7", "多显示器与桌面交互", "支持目标显示器、停靠边缘和水平位置，提供自动收起、悬停展开、点击穿透与鼠标局部避让。"],
            ["8", "设置、缓存与单实例", "统一保存主题、歌词、播放器、布局、屏幕和快捷键设置；提供 LRU 歌词缓存和重复启动复用。"],
        ],
        [700, 2400, 6260],
    )

    doc.add_heading("三、技术实现变化", level=1)
    add_info_table(doc, [
        ("媒体层", "使用统一媒体会话快照、播放器画像、会话选择策略和控制能力模型。"),
        ("歌词层", "组合多个歌词客户端，清洗曲目信息、评分候选、解析同步歌词及已有翻译并持久化缓存。"),
        ("时间轴层", "保存可信采样，使用单调时钟补偿播放位置，处理暂停、恢复、跳转、切歌与过期异步结果。"),
        ("界面层", "使用 WPF 模块宿主、布局投影和交互状态控制器驱动顶部歌词岛及真实跨窗口拖放编辑。"),
        ("数据层", "设置和歌词数据写入 %LOCALAPPDATA%\\LyricsIsland，并提供旧目录迁移与异常回退。"),
    ])

    doc.add_heading("四、与申报材料的对应关系", level=1)
    p = doc.add_paragraph()
    set_run_font(
        p.add_run(
            f"源程序鉴别材料和软件说明书均从当前 {BUILD_VERSION} 工作树生成，"
            f"对外统一登记为 {SOFTWARE_NAME} {VERSION}。源程序样本仅包含 LyricsIsland.App"
            " 和 LyricsIsland.Core 的自有 C#/XAML 代码，不包含第三方依赖库源代码、测试代码和网站代码。"
        ),
        size=10.5,
    )
    add_note(
        doc,
        "提交条件",
        "若本次属于该软件首次登记，登记系统未要求新增版本说明时，可不单独上传本文件；"
        "若此前已有版本登记，或系统／补正通知要求说明高版本变化，则按实际情况核对后提交。",
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    set_run_font(p.add_run(f"著作权人（签名）：{OWNER}"), size=10.5)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    set_run_font(p.add_run("日期：_______年____月____日"), size=10.5)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / f"{SOFTWARE_NAME}{VERSION}-新增版本说明.docx"
    doc.save(path)
    return path


def main() -> None:
    for path in (build_application_brief(), build_version_statement()):
        print(path)


if __name__ == "__main__":
    main()
