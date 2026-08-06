from __future__ import annotations

import re
import zipfile
from pathlib import Path

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT / "output" / "软件著作权登记材料-V2.0.36"
SOURCE = OUTPUT_DIR / "桌面歌词岛软件V2.0.36-源程序鉴别材料.pdf"
MANUAL = OUTPUT_DIR / "桌面歌词岛软件V2.0.36-软件说明书.pdf"
BRIEF = OUTPUT_DIR / "桌面歌词岛软件V2.0.36-申报信息填报底稿.pdf"
VERSION_PDF = OUTPUT_DIR / "桌面歌词岛软件V2.0.36-新增版本说明.pdf"
BRIEF_DOCX = OUTPUT_DIR / "桌面歌词岛软件V2.0.36-申报信息填报底稿.docx"
VERSION_DOCX = OUTPUT_DIR / "桌面歌词岛软件V2.0.36-新增版本说明.docx"


def validate_source() -> None:
    reader = PdfReader(str(SOURCE))
    assert len(reader.pages) == 60, f"source page count={len(reader.pages)}"
    seen_numbers: list[int] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        assert "桌面歌词岛软件" in text, f"missing software name on source page {index}"
        assert f"第 {index} 页 / 共 60 页" in text, f"missing page marker on source page {index}"
        numbers = [int(value) for value in re.findall(r"(?m)^(\d{6})\s*$", text)]
        assert len(numbers) == 50, f"source page {index} has {len(numbers)} numbered lines"
        seen_numbers.extend(numbers)
    assert len(seen_numbers) == 3000
    assert seen_numbers[:1500] == list(range(1, 1501)), "front 1500 lines are not continuous"
    assert all(b == a + 1 for a, b in zip(seen_numbers[1500:], seen_numbers[1501:])), "back 1500 lines are not continuous"


def validate_manual() -> None:
    reader = PdfReader(str(MANUAL))
    assert 8 <= len(reader.pages) < 60, f"manual page count={len(reader.pages)}"
    text = "\n".join((page.extract_text() or "") for page in reader.pages)
    required = [
        "桌面歌词岛软件",
        "V2.0.36",
        "么博丞",
        "独立开发",
        "多播放器识别与选择",
        "歌词检索、匹配与同步",
        "真实歌词岛布局编辑",
        "Apple Music",
        "QQ音乐",
        "网易云音乐",
        "酷狗音乐",
        "酷我音乐",
        "Spotify",
    ]
    for value in required:
        assert value in text, f"manual missing {value}"


def validate_supplemental_pdf(path: Path, expected_pages: int, required: list[str]) -> None:
    reader = PdfReader(str(path))
    assert len(reader.pages) == expected_pages, f"{path.name} pages={len(reader.pages)}"
    text = "\n".join((page.extract_text() or "") for page in reader.pages)
    for value in required:
        assert value in text, f"{path.name} missing {value}"


def validate_docx(path: Path, required: list[str]) -> None:
    with zipfile.ZipFile(path) as archive:
        assert archive.testzip() is None, f"corrupt DOCX: {path.name}"
    doc = Document(path)
    text_parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            text_parts.extend(cell.text for cell in row.cells)
    text = "\n".join(text_parts)
    for value in required:
        assert value in text, f"{path.name} missing {value}"
    section = doc.sections[0]
    assert abs(section.page_width.cm - 21.0) < 0.1
    assert abs(section.page_height.cm - 29.7) < 0.1


def main() -> None:
    for path in (SOURCE, MANUAL, BRIEF, VERSION_PDF):
        assert path.exists() and path.stat().st_size > 50_000, f"missing or too small: {path}"
    for path in (BRIEF_DOCX, VERSION_DOCX):
        assert path.exists() and path.stat().st_size > 20_000, f"missing or too small: {path}"
    validate_source()
    validate_manual()
    validate_supplemental_pdf(
        BRIEF,
        4,
        ["V2.0.36", "必须由申请人确认", "办理入口与规则依据"],
    )
    validate_supplemental_pdf(
        VERSION_PDF,
        2,
        ["V2.0.36", "多播放器媒体会话", "真实界面布局编辑", "著作权人（签名）"],
    )
    validate_docx(BRIEF_DOCX, ["V2.0.36", "必须由申请人确认的事实"])
    validate_docx(VERSION_DOCX, ["V2.0.36", "本版本主要新增与完善内容"])
    print(
        "PASS source_pages=60 manual_pages="
        f"{len(PdfReader(str(MANUAL)).pages)} brief_pages=4 version_pages=2 docx=2"
    )


if __name__ == "__main__":
    main()
